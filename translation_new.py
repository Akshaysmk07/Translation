# streamlit
import streamlit as st
from transformers import MarianMTModel, MarianTokenizer
from docx import Document 
from io import BytesIO

# Mapping of full language names to language codes
language_mapping = {
    "Japanese": "jap",
    "Spanish": "es",
    "German": "de",
    "French": "fr",
    "Italian": "it"
}

st.sidebar.title("Select Language")
selected_language = st.sidebar.selectbox("Select a language:", list(language_mapping.keys()))

st.title("LANGUAGE TRANSLATOR")
uploaded_file = st.file_uploader("Choose a Word document", type=['docx'])

def translate_text(text, target_language):
    # Load pre-trained translation model and tokenizer
    model_name = f'Helsinki-NLP/opus-mt-en-{target_language}'
    model = MarianMTModel.from_pretrained(model_name)
    tokenizer = MarianTokenizer.from_pretrained(model_name)

    # Tokenize input text
    inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True)

    # Translate text
    translated_ids = model.generate(inputs['input_ids'])
    translated_text = tokenizer.batch_decode(translated_ids, skip_special_tokens=True)[0]

    return translated_text

# Function to handle translation from Word file
def handle_translation_from_word(word_file, target_language):
    # Load the input document
    input_doc = Document(word_file)
    
    # Create a new document for translated text
    translated_doc = Document() 
    
    # Iterate through each paragraph in the input document
    for paragraph in input_doc.paragraphs:
        text = paragraph.text
        
        # Translate the paragraph text
        translated_text = translate_text(text, language_mapping[target_language])
        
        # Add the translated text as a new paragraph in the translated document
        translated_doc.add_paragraph(translated_text)

    return translated_doc

if uploaded_file is not None and selected_language:
    if st.button("Translate"):
        translated_document = handle_translation_from_word(uploaded_file, selected_language)
        # Save the translated document to a BytesIO object
        buffer = BytesIO()
        translated_document.save(buffer)
        st.download_button(label='Download Translated Document', data=buffer.getvalue(), file_name='translated_document.docx', mime='application/octet-stream')
elif not selected_language:
    st.error("Please select a language.")
